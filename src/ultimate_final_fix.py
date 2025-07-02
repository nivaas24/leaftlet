"""
Ultimate Final Fix - Direct coordinate copying from original PDF
"""

import fitz
from docx import Document
import re

def get_modifications():
    """Get text modifications from Word document"""
    doc = Document("output/enhanced_modified.docx")
    mappings = {}
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        if "Address:" in text:
            match = re.search(r'Address:\s*(.+)', text)
            if match:
                mappings["123 Sample Street, Example City, EX 12345"] = match.group(1).strip()
        elif "Phone:" in text:
            match = re.search(r'Phone:\s*(.+)', text)
            if match:
                mappings["(555) 123-4567"] = match.group(1).strip()
        elif "Email:" in text:
            match = re.search(r'Email:\s*(.+)', text)
            if match:
                mappings["contact@example.com"] = match.group(1).strip()
    
    # Check for Lorem ipsum change
    modified_text = ' '.join([p.text for p in doc.paragraphs])
    if "Sample text" in modified_text and "Lorem ipsum" not in modified_text:
        mappings["Lorem ipsum"] = "Sample text"
    
    return mappings

def apply_modifications(text, mappings):
    """Apply text modifications"""
    for original, replacement in mappings.items():
        text = text.replace(original, replacement)
    return text

def reconstruct_with_exact_positions():
    """Reconstruct PDF using exact positions from original"""
    
    print("=== ULTIMATE FINAL FIX ===")
    
    # Get modifications
    mappings = get_modifications()
    print(f"Applying {len(mappings)} modifications:")
    for orig, new in mappings.items():
        print(f"  '{orig}' -> '{new}'")
    
    # Open original PDF
    original_doc = fitz.open("samples/enhanced_sample.pdf")
    original_page = original_doc[0]
    
    # Create new PDF with exact same page size
    new_doc = fitz.open()
    new_page = new_doc.new_page(width=original_page.rect.width, height=original_page.rect.height)
    
    print(f"\nPage size: {original_page.rect.width} x {original_page.rect.height}")
    
    # Get all text blocks from original
    text_blocks = original_page.get_text("dict")
    
    print(f"Processing {len(text_blocks['blocks'])} blocks...")
    
    # Process each block, line, and span
    for block_idx, block in enumerate(text_blocks["blocks"]):
        if "lines" in block:
            print(f"\nBlock {block_idx}:")
            
            for line_idx, line in enumerate(block["lines"]):
                for span_idx, span in enumerate(line["spans"]):
                    original_text = span['text']
                    if not original_text.strip():
                        continue
                    
                    # Apply modifications
                    modified_text = apply_modifications(original_text, mappings)
                    
                    # Use EXACT original positioning
                    rect = fitz.Rect(span['bbox'])
                    
                    # Get original font info
                    font_name = span['font']
                    font_size = span['size']
                    color = span['color']
                    
                    # Map fonts to PyMuPDF fonts
                    if 'Bold' in font_name:
                        font = "helvetica-bold"
                    elif 'Oblique' in font_name:
                        font = "helvetica-oblique" 
                    else:
                        font = "helvetica"
                    
                    # Convert color
                    if color == 139:  # Dark blue
                        text_color = (0, 0, 139/255)
                    else:
                        text_color = (0, 0, 0)  # Black
                    
                    # Insert text at EXACT original position
                    new_page.insert_text(
                        rect.tl,  # Top-left point - EXACT original position
                        modified_text,
                        fontname=font,
                        fontsize=font_size,
                        color=text_color
                    )
                    
                    if original_text != modified_text:
                        print(f"  Modified: '{original_text[:20]}...' -> '{modified_text[:20]}...'")
                    elif original_text in ["Enhanced Sample Document", "Sample Image 1", "Sample Image 2"]:
                        print(f"  Key element: '{original_text}' at exact position {rect.tl}")
    
    # Copy images from original - INSERT AT ALL RECTANGLE POSITIONS
    print("\nCopying images...")
    image_list = original_page.get_images(full=True)
    for img_idx, img in enumerate(image_list):
        # Get image data from original
        image_data = original_doc.extract_image(img[0])
        image_bytes = image_data["image"]
        
        # Get ALL image rectangles from original (one image can appear in multiple positions)
        img_rects = original_page.get_image_rects(img[0])
        print(f"  Image {img_idx + 1} found at {len(img_rects)} position(s)")
        
        # Insert image at EVERY position it appears in the original
        for rect_idx, img_rect in enumerate(img_rects):
            new_page.insert_image(img_rect, stream=image_bytes)
            print(f"    Position {rect_idx + 1}: {img_rect}")
    
    # Save the new PDF
    new_doc.save("output/ULTIMATE_FINAL_FIX.pdf")
    
    # Clean up
    original_doc.close()
    new_doc.close()
    
    print(f"\n✅ ULTIMATE FIX COMPLETE: output/ULTIMATE_FINAL_FIX.pdf")
    print("\n🎯 This uses EXACT coordinate copying from the original PDF")
    print("   Every element positioned at precisely the same coordinates")

if __name__ == "__main__":
    reconstruct_with_exact_positions()
